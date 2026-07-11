"""
Document Loader for RAG.
Loads and processes course documents: FAQ.pdf, course.txt, RTFM.docx.
"""

from pathlib import Path
from typing import List

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from config import DOCUMENTS_DIR, RAG_CHUNK_SIZE, RAG_CHUNK_OVERLAP, SUPPORTED_DOC_EXTENSIONS
from utils.logging import logger


class DocumentLoader:
    """Loads and processes documents for RAG."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=RAG_CHUNK_SIZE,
            chunk_overlap=RAG_CHUNK_OVERLAP,
            length_function=len,
        )

    def _load_docx(self, file_path: Path) -> List[Document]:
        """Load DOCX file using python-docx."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        if not text:
            raise ValueError(f"DOCX file is empty: {file_path.name}")

        return [Document(page_content=text, metadata={"source": file_path.name})]

    def load_document(self, file_path: Path) -> List:
        """Load a single document and split into chunks."""
        try:
            file_path = Path(file_path)
            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                loader = PyPDFLoader(str(file_path))
                documents = loader.load()
            elif suffix in (".txt", ".md"):
                loader = TextLoader(str(file_path), encoding="utf-8")
                documents = loader.load()
            elif suffix == ".docx":
                documents = self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")

            chunks = self.text_splitter.split_documents(documents)

            for chunk in chunks:
                chunk.metadata["source"] = file_path.name
                chunk.metadata["file_path"] = str(file_path)

            logger.info(f"Loaded {len(chunks)} chunks from {file_path.name}")
            return chunks

        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            raise

    def load_directory(self, directory: Path = DOCUMENTS_DIR) -> List:
        """Load all supported documents from a directory."""
        try:
            directory = Path(directory)
            all_chunks = []

            for file_path in directory.rglob("*"):
                if file_path.suffix.lower() in SUPPORTED_DOC_EXTENSIONS:
                    try:
                        chunks = self.load_document(file_path)
                        all_chunks.extend(chunks)
                    except Exception as e:
                        logger.warning(f"Skipping {file_path.name}: {e}")

            logger.info(f"Loaded {len(all_chunks)} total chunks from {directory}")
            return all_chunks

        except Exception as e:
            logger.error(f"Error loading directory {directory}: {e}")
            raise

    def load_text(self, text: str, source: str = "manual_input") -> List:
        """Load text directly and split into chunks."""
        document = Document(page_content=text, metadata={"source": source})
        chunks = self.text_splitter.split_documents([document])
        logger.info(f"Created {len(chunks)} chunks from text input")
        return chunks


document_loader = DocumentLoader()
